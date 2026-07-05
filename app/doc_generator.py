"""
Renders the agent's plan + generated section text into a polished .docx
using python-docx. Kept deliberately simple (no external template) so the
whole project runs with nothing but `pip install -r requirements.txt`.
"""
import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.config import settings
from app.models import ExecutionPlan

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)


def _add_title_page(doc: Document, plan: ExecutionPlan, request_text: str):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(plan.title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(plan.document_type.replace("_", " ").title())
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')} · Autonomous Agent Draft")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    src = doc.add_paragraph()
    src_run = src.add_run("Original request: ")
    src_run.bold = True
    src.add_run(request_text)

    if plan.assumptions:
        doc.add_paragraph()
        a_head = doc.add_paragraph()
        a_head.add_run("Assumptions made by the agent:").bold = True
        for a in plan.assumptions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(a)

    doc.add_page_break()


def _looks_like_bullets(text: str) -> bool:
    lines = [l for l in text.splitlines() if l.strip()]
    if not lines:
        return False
    bullet_lines = [l for l in lines if l.strip().startswith(("-", "*", "•"))]
    return len(bullet_lines) >= max(2, len(lines) // 2)


def _add_section(doc: Document, heading: str, content: str):
    doc.add_heading(heading, level=1)

    if _looks_like_bullets(content):
        for line in content.splitlines():
            line = line.strip()
            if not line:
                continue
            cleaned = re.sub(r"^[-*•]\s*", "", line)
            doc.add_paragraph(cleaned, style="List Bullet")
    else:
        for para in content.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)


def build_document(plan: ExecutionPlan, section_content: dict, request_text: str) -> str:
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    _add_title_page(doc, plan, request_text)

    for i, sec_name in enumerate(plan.sections, start=1):
        content = section_content.get(sec_name, "")
        _add_section(doc, f"{i}. {sec_name}", content)

    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", plan.title)[:60].strip("_") or "document"
    filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(settings.OUTPUT_DIR, filename)
    doc.save(filepath)
    return filename
